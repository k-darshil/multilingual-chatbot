"""
Document processing module for extracting text from various file formats.
Supports PDF, DOCX, TXT files with fallback options.
"""

import io
import tempfile
import os
from typing import Optional, Dict, Any, List
# import streamlit as st  # Removed for Gradio compatibility
import PyPDF2
import pdfplumber
from docx import Document
import chardet
from pathlib import Path

try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    print("DocLing not available. Using fallback PDF extraction methods.")

class DocumentProcessor:
    """Class for processing and extracting text from various document formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.doc']
        self.converter = None
        
        # Initialize DocLing if available
        if DOCLING_AVAILABLE:
            try:
                self.converter = DocumentConverter()
            except Exception as e:
                print(f"Could not initialize DocLing: {e}")
                self.converter = None
    
    def process_file(self, uploaded_file, target_language: str = 'en', translation_service=None) -> Dict[str, Any]:
        """
        Process an uploaded file and extract text content with language detection and translation.
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            target_language: Target language code for translation (default: 'en')
            translation_service: TranslationService instance for language operations
            
        Returns:
            Dict containing extracted text (original and translated), metadata, and processing info
        """
        if uploaded_file is None:
            return {"success": False, "error": "No file uploaded"}
        
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        if file_extension not in self.supported_formats:
            return {
                "success": False, 
                "error": f"Unsupported file format: {file_extension}. Supported formats: {', '.join(self.supported_formats)}"
            }
        
        try:
            # Read file content
            file_content = uploaded_file.read()
            
            # Reset file pointer for potential re-reading
            uploaded_file.seek(0)
            
            # Extract text based on file type
            if file_extension == '.pdf':
                extracted_text = self._extract_pdf_text(file_content, uploaded_file.name)
            elif file_extension in ['.docx', '.doc']:
                extracted_text = self._extract_docx_text(file_content)
            elif file_extension == '.txt':
                extracted_text = self._extract_txt_text(file_content)
            else:
                return {"success": False, "error": f"Handler not implemented for {file_extension}"}
            
            if not extracted_text.strip():
                return {"success": False, "error": "No text could be extracted from the document"}
            
            # Language detection and translation
            detected_language = None
            translated_text = extracted_text
            translation_needed = False
            translation_success = True
            translation_error = None
            
            if translation_service:
                try:
                    # Detect language of extracted text
                    detected_language = translation_service.detect_language(extracted_text)
                    print(f"🔍 Detected document language: {translation_service.get_language_name(detected_language) if detected_language else 'Unknown'}")
                    
                    # Check if translation is needed
                    if detected_language and detected_language != target_language:
                        translation_needed = True
                        print(f"🌍 Translating document from {translation_service.get_language_name(detected_language)} to {translation_service.get_language_name(target_language)}")
                        
                        print("Translating document from ", detected_language, "to", target_language)
                        # Translate the text
                        print("Translating document...")
                        translation_result = translation_service.translate_text(
                            extracted_text, 
                            target_language, 
                            detected_language,
                            uploaded_file.name
                        )
                        
                        # print("Translation result", translation_result)
                        if translation_result["success"]:
                            translated_text = translation_result["translated_text"]
                            print(f"✅ Document translated successfully using {translation_result.get('method', 'unknown method')}")
                        else:
                            translation_success = False
                            translation_error = translation_result.get("error", "Translation failed")
                            print(f"⚠️ Translation failed: {translation_error}. Using original text.")
                            translated_text = extracted_text
                    else:
                        print(f"ℹ️ No translation needed - document is already in {translation_service.get_language_name(target_language)}")
                        
                except Exception as e:
                    print(f"⚠️ Language processing failed: {e}. Using original text.")
                    translation_error = str(e)
                    translation_success = False
            
            # Calculate metadata for the final text (translated if successful, original otherwise)
            final_text = translated_text
            word_count = len(final_text.split())
            char_count = len(final_text)
            
            return {
                "success": True,
                "text": final_text,  # This will be the translated text for indexing
                "original_text": extracted_text,  # Keep original for reference
                "translated_text": translated_text if translation_needed else None,
                "filename": uploaded_file.name,
                "file_size": len(file_content),
                "file_type": file_extension,
                "word_count": word_count,
                "char_count": char_count,
                "pages": self._estimate_pages(final_text),
                "detected_language": detected_language,
                "target_language": target_language,
                "translation_needed": translation_needed,
                "translation_success": translation_success,
                "translation_error": translation_error
            }
            
        except Exception as e:
            return {"success": False, "error": f"Error processing file: {str(e)}"}
    
    def _extract_pdf_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF using multiple methods."""
        
        print("Docling available: ", DOCLING_AVAILABLE)
        
        
        # Try DocLing first if available
        if self.converter and DOCLING_AVAILABLE:
            print("using docling")
            try:
                # Save content to temporary file for DocLing
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(file_content)
                    tmp_file_path = tmp_file.name
                
                # Convert with DocLing
                result = self.converter.convert(tmp_file_path)
                print("Docling conversion successful")
                
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
                # Extract text from DocLing result
                if hasattr(result, 'document') and hasattr(result.document, 'export_to_text'):
                    exported_text = result.document.export_to_text()
                    # print("Docling exported_text: ", exported_text)
                    return exported_text
                
            except Exception as e:
                print(f"DocLing extraction failed: {e}. Trying fallback methods.")
        
        # Fallback to pdfplumber
        try:
            with io.BytesIO(file_content) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    
                    if text_parts:
                        return '\n\n'.join(text_parts)
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}. Trying PyPDF2.")
        
        # Fallback to PyPDF2
        try:
            with io.BytesIO(file_content) as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    return '\n\n'.join(text_parts)
                
        except Exception as e:
            raise Exception(f"All PDF extraction methods failed. Last error: {e}")
        
        raise Exception("Could not extract text from PDF file")
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            with io.BytesIO(file_content) as docx_file:
                doc = Document(docx_file)
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(' | '.join(row_text))
                
                return '\n\n'.join(text_parts)
                
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {e}")
    
    def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT file with encoding detection."""
        try:
            # Detect encoding
            detected = chardet.detect(file_content)
            encoding = detected.get('encoding', 'utf-8')
            
            # Decode with detected encoding
            text = file_content.decode(encoding)
            return text
            
        except Exception as e:
            # Fallback to common encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding)
                except:
                    continue
            
            raise Exception(f"Failed to decode text file: {e}")
    
    def _estimate_pages(self, text: str) -> int:
        """Estimate number of pages based on text length."""
        # Rough estimation: 250 words per page
        words = len(text.split())
        return max(1, words // 250)
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for better RAG performance.
        
        Args:
            text: Input text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If not the last chunk, try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                for i in range(end - 100, end):
                    if i > start and text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = max(start + 1, end - overlap)
            
            # Avoid infinite loops
            if start >= len(text):
                break
        
        return chunks
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return self.supported_formats.copy()
    
    def validate_file(self, uploaded_file) -> tuple[bool, str]:
        """
        Validate uploaded file.
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        if uploaded_file is None:
            return False, "No file uploaded"
        
        # Check file extension
        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in self.supported_formats:
            return False, f"Unsupported file format: {file_extension}"
        
        # Check file size (convert MB to bytes)
        max_size_bytes = 50 * 1024 * 1024  # 50 MB
        if uploaded_file.size > max_size_bytes:
            return False, f"File size ({uploaded_file.size / 1024 / 1024:.1f} MB) exceeds maximum allowed size (50 MB)"
        
        return True, "" 